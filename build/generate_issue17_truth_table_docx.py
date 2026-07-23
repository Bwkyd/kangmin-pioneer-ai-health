#!/usr/bin/env python3
"""Generate the concise Issue #17 customer confirmation DOCX."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs/客户确认/Issue-17-证型32组真值表-待确认.docx"
SKILL_SCRIPTS = Path(
    "/Users/chenqiqiang/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.715.12143/skills/documents/scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "18324A"
MUTED = "5F6B76"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CAUTION = "FFF4CE"
BLACK = "000000"
LATIN_FONT = "Calibri"
CJK_FONT = "PingFang SC"
CONTENT_WIDTH = 9360
CELL_MARGINS = {"top": 100, "bottom": 100, "start": 120, "end": 120}


def set_run(run, *, size=11, bold=False, color=BLACK) -> None:
    run.font.name = LATIN_FONT
    r_fonts = run._element.get_or_add_rPr().rFonts
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), CJK_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph(
    paragraph,
    *,
    before=0,
    after=6,
    line=1.2,
    keep=False,
    alignment=None,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep
    if alignment is not None:
        paragraph.alignment = alignment


def add_text(doc, text="", *, bold=False, color=BLACK, after=6, keep=False):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, after=after, keep=keep)
    run = paragraph.add_run(text)
    set_run(run, bold=bold, color=color)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    tokens = {
        1: (16, BLUE, 16, 8),
        2: (13, BLUE, 11, 5),
    }[level]
    run = paragraph.add_run(text)
    set_run(run, size=tokens[0], bold=True, color=tokens[1])
    set_paragraph(
        paragraph,
        before=tokens[2],
        after=tokens[3],
        line=1.0,
        keep=True,
    )
    return paragraph


def add_checkbox(doc, text, *, after=4):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, after=after)
    box = paragraph.add_run("□ ")
    set_run(box, size=12, bold=True, color=DARK_BLUE)
    label = paragraph.add_run(text)
    set_run(label)
    return paragraph


def set_shading(cell, fill) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_borders(cell, color="C9D1D9") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)
        borders.append(element)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def style_table(table, *, font_size=10, header=True, center_columns=None) -> None:
    center_columns = center_columns or set()
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_borders(cell)
            if header and row_index == 0:
                set_shading(cell, LIGHT_BLUE)
            for paragraph in cell.paragraphs:
                set_paragraph(paragraph, after=0, line=1.1)
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if column_index in center_columns
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in paragraph.runs:
                    set_run(
                        run,
                        size=font_size,
                        bold=header and row_index == 0,
                        color=NAVY if header and row_index == 0 else BLACK,
                    )
    if header:
        repeat_header(table.rows[0])


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_shading(cell, CAUTION)
    set_borders(cell, "E3C75F")
    paragraph = cell.paragraphs[0]
    set_paragraph(paragraph, after=0)
    run = paragraph.add_run(text)
    set_run(run, size=10.5, bold=True, color="7A5A00")
    apply_table_geometry(
        table,
        [CONTENT_WIDTH],
        table_width_dxa=CONTENT_WIDTH,
        indent_dxa=160,
        cell_margins_dxa={"top": 120, "bottom": 120, "start": 160, "end": 160},
    )
    add_text(doc, "", after=1)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    set_run(run, size=9, color=MUTED)


def configure_document(doc) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for level, size, color, before, after in (
        (1, 16, BLUE, 16, 8),
        (2, 13, BLUE, 11, 5),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.clear()
    set_paragraph(header, after=0, line=1.0)
    run = header.add_run("抗敏先锋 AI 鼻健康管理系统　|　规则确认单")
    set_run(run, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.clear()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(footer, after=0, line=1.0)
    prefix = footer.add_run("待确认　·　第 ")
    set_run(prefix, size=9, color=MUTED)
    add_page_field(footer)
    suffix = footer.add_run(" 页")
    set_run(suffix, size=9, color=MUTED)


def build_document():
    doc = Document()
    configure_document(doc)

    kicker = doc.add_paragraph()
    set_paragraph(kicker, after=2)
    set_run(kicker.add_run("客户／导师确认文件"), size=10, bold=True, color=BLUE)
    title = doc.add_paragraph()
    set_paragraph(title, after=5, line=1.0)
    set_run(title.add_run("过敏性鼻炎辨证规则关键确认单"), size=22, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    set_paragraph(subtitle, after=12)
    set_run(
        subtitle.add_run("请只确认以下关键问题；完整 32 组规则表由开发方自动生成。"),
        size=11.5,
        color=MUTED,
    )
    add_note(
        doc,
        "收到确认后，开发方才会把这些决定写入正式规则和自动化测试。",
    )

    add_heading(doc, "一、请确认已有结论", 1)
    add_heading(doc, "1. 一期使用 5 个证型", 2)
    syndromes = [
        "肺经伏热，上犯鼻窍",
        "肺气虚寒，卫表不固",
        "脾气虚弱，清阳不升",
        "肾阳不足，温煦失职",
        "寒热错杂，虚实并见",
    ]
    syndrome_table = doc.add_table(rows=1, cols=2)
    syndrome_table.rows[0].cells[0].text = "序号"
    syndrome_table.rows[0].cells[1].text = "证型"
    for index, syndrome in enumerate(syndromes, 1):
        cells = syndrome_table.add_row().cells
        cells[0].text = str(index)
        cells[1].text = syndrome
    apply_table_geometry(
        syndrome_table,
        [850, 8510],
        table_width_dxa=CONTENT_WIDTH,
        indent_dxa=120,
        cell_margins_dxa=CELL_MARGINS,
    )
    style_table(syndrome_table, center_columns={0})
    add_checkbox(doc, "确认")
    add_checkbox(doc, "需要修改：________________________________________")

    add_heading(doc, "2. 冲突时优先判“寒热错杂”", 2)
    add_text(
        doc,
        "当用户“口渴、手脚冷、但不乏力”时，判为“寒热错杂”，不判“肺经伏热”。",
    )
    add_checkbox(doc, "确认")
    add_checkbox(doc, "需要修改：________________________________________")

    add_heading(doc, "3. 严重程度使用“轻／中／重”三级", 2)
    add_checkbox(doc, "确认")
    add_text(doc, "请补充三级的明确分界：", bold=True, keep=True)
    severity = doc.add_table(rows=1, cols=2)
    severity.rows[0].cells[0].text = "等级"
    severity.rows[0].cells[1].text = "明确判断条件"
    for level in ("轻度", "中度", "重度"):
        cells = severity.add_row().cells
        cells[0].text = level
        cells[1].text = "\n"
    apply_table_geometry(
        severity,
        [1300, 8060],
        table_width_dxa=CONTENT_WIDTH,
        indent_dxa=140,
        cell_margins_dxa={"top": 120, "bottom": 120, "start": 140, "end": 140},
    )
    style_table(severity, center_columns={0})

    add_heading(doc, "二、以下 4 种情况应该怎么处理？", 1)
    add_text(
        doc,
        "请为每种情况填写：直接判为什么证型、继续补问什么，或者显示“信息不足”。",
    )
    cases = [
        ("1", "不口渴、不乏力、手脚不冷、不怕风，也不特别怕冷。"),
        ("2", "不口渴、不乏力、手脚冷、不怕风，但不觉得特别怕冷。"),
        ("3", "不口渴、不乏力、手脚冷、怕风，但不觉得特别怕冷。"),
        ("4", "不口渴、不乏力、手脚不冷、不怕风，但觉得特别怕冷。"),
    ]
    cases_table = doc.add_table(rows=1, cols=3)
    for cell, text in zip(cases_table.rows[0].cells, ("情况", "回答组合", "处理方式")):
        cell.text = text
    for number, description in cases:
        cells = cases_table.add_row().cells
        cells[0].text = number
        cells[1].text = description
        cells[2].text = "\n\n"
        set_shading(cells[2], LIGHT_GRAY)
    apply_table_geometry(
        cases_table,
        [800, 5260, 3300],
        table_width_dxa=CONTENT_WIDTH,
        indent_dxa=120,
        cell_margins_dxa={"top": 130, "bottom": 130, "start": 120, "end": 120},
    )
    style_table(cases_table, font_size=9.8, center_columns={0})

    add_heading(doc, "三、“不确定”答案", 1)
    add_text(doc, "如果用户对关键问题回答“不确定”，是否按以下方式处理：")
    add_checkbox(doc, "不把“不确定”当成“否”。")
    add_checkbox(doc, "继续补问相关症状。")
    add_checkbox(doc, "仍不能确认时显示“信息不足”，不输出证型方案。")
    add_checkbox(doc, "以上处理方式确认")
    add_checkbox(doc, "需要修改：________________________________________")

    add_heading(doc, "四、确认信息", 1)
    approval = doc.add_table(rows=3, cols=2)
    values = [
        ("确认人", ""),
        ("角色", "□ 导师　□ 临床负责人　□ 其他：____________"),
        ("确认日期", "________年____月____日"),
    ]
    for row, (label, value) in zip(approval.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
        set_shading(row.cells[0], LIGHT_BLUE)
    apply_table_geometry(
        approval,
        [1700, 7660],
        table_width_dxa=CONTENT_WIDTH,
        indent_dxa=140,
        cell_margins_dxa={"top": 140, "bottom": 140, "start": 140, "end": 140},
    )
    style_table(approval, header=False, center_columns={0})

    doc.core_properties.title = "过敏性鼻炎辨证规则关键确认单"
    doc.core_properties.subject = "Issue #17 客户导师或临床负责人确认"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    return doc


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
